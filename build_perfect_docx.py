import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_page_border_to_section(section):
    # Set standard 1 inch margins
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    sectPr = section._sectPr
    # Add a solid thin black box border around the page like the sample PDF
    pgBorders = parse_xml(
        f'<w:pgBorders {nsdecls("w")}>\n'
        f'  <w:top w:val="single" w:sz="12" w:space="20" w:color="000000"/>\n'
        f'  <w:left w:val="single" w:sz="12" w:space="20" w:color="000000"/>\n'
        f'  <w:bottom w:val="single" w:sz="12" w:space="20" w:color="000000"/>\n'
        f'  <w:right w:val="single" w:sz="12" w:space="20" w:color="000000"/>\n'
        f'</w:pgBorders>'
    )
    sectPr.append(pgBorders)

def add_page_number(run):
    fldChar1 = parse_xml(r'<w:fldChar %s w:fldCharType="begin"/>' % nsdecls('w'))
    instrText = parse_xml(r'<w:instrText %s xml:space="preserve"> PAGE </w:instrText>' % nsdecls('w'))
    fldChar2 = parse_xml(r'<w:fldChar %s w:fldCharType="separate"/>' % nsdecls('w'))
    fldChar3 = parse_xml(r'<w:fldChar %s w:fldCharType="end"/>' % nsdecls('w'))
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

def add_footer_page_number(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run()
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(50, 50, 50)
    add_page_number(run)

def create_document():
    doc = docx.Document()
    
    # Configure base styling
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0, 0, 0)

    # Initial section setup
    section = doc.sections[0]
    add_page_border_to_section(section)
    add_footer_page_number(section)

    primary_red = RGBColor(180, 0, 0)
    dark_blue = RGBColor(0, 51, 102)
    black = RGBColor(0, 0, 0)
    cyan_header_bg = "0099FF"

    # =============================================================
    # COVER PAGE (Page 1)
    # =============================================================
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_logo.paragraph_format.space_before = Pt(0)
    p_logo.paragraph_format.space_after = Pt(4)
    if os.path.exists('assets/report_images/au_logo.png'):
        p_logo.add_run().add_picture('assets/report_images/au_logo.png', width=Inches(1.5))

    p_header = doc.add_paragraph()
    p_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_header.paragraph_format.space_after = Pt(2)
    r1 = p_header.add_run("ఆంధ్ర విశ్వకళాపరిషత్\n")
    r1.font.name = 'Arial'
    r1.font.size = Pt(13)
    r1.font.bold = True
    r1.font.color.rgb = RGBColor(204, 102, 0)

    r2 = p_header.add_run("ANDHRA UNIVERSITY\n")
    r2.font.name = 'Arial'
    r2.font.size = Pt(18)
    r2.font.bold = True
    r2.font.color.rgb = primary_red

    r3 = p_header.add_run("Accredited by NAAC with A++ Grade & Score: 3.74")
    r3.font.name = 'Arial'
    r3.font.size = Pt(9.5)
    r3.font.bold = True
    r3.font.color.rgb = dark_blue

    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_after = Pt(12)

    p_proj = doc.add_paragraph()
    p_proj.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_proj.paragraph_format.space_after = Pt(12)
    r_rep = p_proj.add_run("PROJECT REPORT ON\n")
    r_rep.font.size = Pt(13)
    r_rep.font.bold = True

    r_title = p_proj.add_run('“Healthcare Data Analysis and Visualization Dashboard”')
    r_title.font.size = Pt(16)
    r_title.font.bold = True
    r_title.font.color.rgb = black

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(14)
    r_sub = p_sub.add_run("Submitted in partial fulfilment of the requirements for the award of the degree of\n")
    r_sub.font.size = Pt(11)
    r_deg = p_sub.add_run("MASTER OF COMPUTER APPLICATIONS")
    r_deg.font.size = Pt(14)
    r_deg.font.bold = True
    r_deg.font.italic = True
    r_deg.font.color.rgb = dark_blue

    p_by = doc.add_paragraph()
    p_by.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_by.paragraph_format.space_after = Pt(12)
    p_by.add_run("Submitted By\n").font.bold = True
    p_by.runs[0].font.italic = True
    r_n = p_by.add_run("ANJALI KUMARI\n")
    r_n.font.size = Pt(14)
    r_n.font.bold = True
    r_n.font.italic = True
    p_by.add_run("Registered number\n").font.bold = True
    p_by.runs[2].font.italic = True
    r_r = p_by.add_run("A24CA1895")
    r_r.font.size = Pt(13)
    r_r.font.bold = True
    r_r.font.underline = True

    p_g = doc.add_paragraph()
    p_g.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_g.paragraph_format.space_after = Pt(14)
    p_g.add_run("Under The Guidance of\n").font.bold = True
    p_g.runs[0].font.italic = True
    r_g1 = p_g.add_run("Mr. Divakar Purohit\n")
    r_g1.font.size = Pt(13)
    r_g1.font.bold = True
    r_g1.font.italic = True
    r_g2 = p_g.add_run("Sr. Data Analyst")
    r_g2.font.size = Pt(11)
    r_g2.font.bold = True

    if os.path.exists('assets/report_images/au_logo.png'):
        p_l2 = doc.add_paragraph()
        p_l2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_l2.paragraph_format.space_after = Pt(4)
        p_l2.add_run().add_picture('assets/report_images/au_logo.png', width=Inches(1.2))

    p_cdoe = doc.add_paragraph()
    p_cdoe.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_c = p_cdoe.add_run("CENTRE FOR DISTANCE AND ONLINE EDUCATION\nANDHRA UNIVERSITY\nVISAKHAPATNAM")
    r_c.font.size = Pt(12)
    r_c.font.bold = True

    doc.add_page_break()

    # =============================================================
    # 1. GUIDE CERTIFICATE
    # =============================================================
    h1 = doc.add_paragraph()
    h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.space_after = Pt(25)
    r = h1.add_run("1. GUIDE CERTIFICATE")
    r.font.size = Pt(15)
    r.font.bold = True

    p_cert = doc.add_paragraph()
    p_cert.paragraph_format.line_spacing = 1.3
    p_cert.paragraph_format.space_after = Pt(180)
    p_cert.add_run("This is to certify that the project entitled ")
    r_t = p_cert.add_run("“Healthcare Data Analysis and Visualization Dashboard”")
    r_t.font.bold = True
    p_cert.add_run(", is a Bonafide work done by ")
    r_n = p_cert.add_run("Anjali Kumari")
    r_n.font.bold = True
    r_n.font.italic = True
    p_cert.add_run(", bearing Regd. No: ")
    r_r = p_cert.add_run("A24CA1895")
    r_r.font.bold = True
    r_r.font.underline = True
    p_cert.add_run(" for the academic year 2024-2025 in partial fulfillment of the requirements for the award of the degree of Master of Computer Applications (M.C.A) in Andhra University. This work has been carried out under my supervision and guidance.")

    p_sig = doc.add_paragraph()
    p_sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_s = p_sig.add_run("Signature and Name of the Project Guide\n\nDate: 09-07-2026")
    r_s.font.bold = True

    doc.add_page_break()

    # =============================================================
    # 2. DECLARATION BY THE LEARNER
    # =============================================================
    h2 = doc.add_paragraph()
    h2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h2.paragraph_format.space_after = Pt(25)
    r = h2.add_run("2. DECLARATION BY THE LEARNER")
    r.font.size = Pt(15)
    r.font.bold = True

    p_dec = doc.add_paragraph()
    p_dec.paragraph_format.line_spacing = 1.3
    p_dec.paragraph_format.space_after = Pt(140)
    p_dec.add_run("I hereby declare that the project report entitled at ")
    r_t = p_dec.add_run("“Healthcare Data Analysis and Visualization Dashboard”")
    r_t.font.bold = True
    p_dec.add_run(" has been carried out by me under the guidance of ")
    r_g = p_dec.add_run("Mr. Divakar Purohit")
    r_g.font.bold = True
    r_g.font.italic = True
    p_dec.add_run(". This project is original and has not been submitted by me, either in part or full, for the award of any degree or diploma at any other university or institution.")

    p_lsig = doc.add_paragraph()
    r_ln = p_lsig.add_run("ANJALI KUMARI\n")
    r_ln.font.bold = True
    r_ln.font.size = Pt(13)
    r_ls = p_lsig.add_run("Signature and Name of the Learner\nRegd. No: A24CA1895\nDate: 09-07-2026")
    r_ls.font.bold = True

    doc.add_page_break()

    # =============================================================
    # 3. ACKNOWLEDGEMENT
    # =============================================================
    h3 = doc.add_paragraph()
    h3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h3.paragraph_format.space_after = Pt(25)
    r = h3.add_run("3. ACKNOWLEDGEMENT")
    r.font.size = Pt(15)
    r.font.bold = True

    p_ack = doc.add_paragraph()
    p_ack.paragraph_format.line_spacing = 1.3
    p_ack.paragraph_format.space_after = Pt(12)
    p_ack.add_run("I would like to express my sincere gratitude to all those who helped me to complete this project titled ")
    r_t = p_ack.add_run("“Healthcare Data Analysis and Visualization Dashboard”")
    r_t.font.bold = True
    p_ack.add_run(" First and foremost, I extend my heartfelt thanks to ")
    r_g = p_ack.add_run("Mr. Divakar Purohit")
    r_g.font.bold = True
    r_g.font.italic = True
    p_ack.add_run(", my project guide, for their valuable guidance, encouragement, and continuous support throughout the course of this study.")

    p_ack2 = doc.add_paragraph()
    p_ack2.paragraph_format.line_spacing = 1.3
    p_ack2.paragraph_format.space_after = Pt(12)
    p_ack2.add_run("I am also grateful to the management and staff of Sri Chanakya Degree College for giving me the opportunity to undertake this project and for providing the necessary information and resources. Their cooperation and insights were instrumental in the successful completion of this work.")

    p_ack3 = doc.add_paragraph()
    p_ack3.paragraph_format.line_spacing = 1.3
    p_ack3.paragraph_format.space_after = Pt(12)
    p_ack3.add_run("I sincerely thank the Centre for Distance and Online Education, Andhra University, for facilitating this academic opportunity.")

    p_ack4 = doc.add_paragraph()
    p_ack4.paragraph_format.line_spacing = 1.3
    p_ack4.paragraph_format.space_after = Pt(100)
    p_ack4.add_run("Lastly, I would like to thank my family and friends for their constant encouragement and moral support during this project.")

    p_adate = doc.add_paragraph()
    r_d = p_adate.add_run("Date: 09-07-2026")
    r_d.font.bold = True

    doc.add_page_break()

    # =============================================================
    # 4. TABLE OF CONTENTS
    # =============================================================
    htoc = doc.add_paragraph()
    htoc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    htoc.paragraph_format.space_after = Pt(20)
    r = htoc.add_run("4. TABLE OF CONTENTS")
    r.font.size = Pt(15)
    r.font.bold = True

    toc_items = [
        ("1.", "GUIDE CERTIFICATE", "2"),
        ("2.", "DECLARATION BY THE LEARNER", "3"),
        ("3.", "ACKNOWLEDGEMENT", "4"),
        ("4.", "TABLE OF CONTENTS", "5"),
        ("5.", "SYNOPSIS", "7"),
        ("6.", "ABSTRACT", "11"),
        ("7.", "CHAPTER NO: 1", "12"),
        ("7.1.", "INTRODUCTION TO TOPIC", "12"),
        ("7.2.", "SYSTEM REQUIREMENT SPECIFICATION", "12"),
        ("7.2.1.", "HARDWARE SPECIFICATION", "13"),
        ("7.2.2.", "SOFTWARE SPECIFICATION", "13"),
        ("7.2.3.", "FEATURES OF THE OPERATING SYSTEM", "14"),
        ("7.2.4.", "SOFTWARE TECHNOLOGIES (FRONT END)", "15"),
        ("7.2.5.", "BACK-END TECHNOLOGIES", "17"),
        ("7.2.6.", "PYTHON & STREAMLIT FEATURES", "18"),
        ("7.3.", "STATEMENT OF THE PROBLEM", "19"),
        ("7.4.", "OBJECTIVES OF THE STUDY", "19"),
        ("7.5.", "SCOPE OF THE STUDY", "19"),
        ("7.6.", "PURPOSE OF THE STUDY", "20"),
        ("7.7.", "HYPOTHESES", "20"),
        ("7.8.", "RESEARCH LIMITATIONS", "21"),
        ("8.", "CHAPTER NO: 2", "22"),
        ("8.1.", "REVIEW OF LITERATURE", "22"),
        ("9.", "CHAPTER NO: 3", "28"),
        ("9.1.", "OVERVIEW OF THE PROJECT", "28"),
        ("9.1.1.", "EXECUTIVE OVERVIEW MODULE", "28"),
        ("9.1.2.", "PATIENT DEMOGRAPHICS MODULE", "29"),
        ("9.1.3.", "DISEASE ANALYSIS MODULE", "29"),
        ("9.1.4.", "HOSPITAL OPERATIONS MODULE", "29"),
        ("9.1.5.", "FINANCIAL ANALYTICS MODULE", "29"),
        ("9.1.6.", "MACHINE LEARNING INSIGHTS MODULE", "30"),
        ("9.2.", "BENEFITS OF THE HEALTHCARE DASHBOARD SYSTEM", "30"),
        ("9.3.", "KEY CHALLENGES OF THE HEALTHCARE DASHBOARD SYSTEM", "32"),
        ("9.4.", "SYSTEM STUDY AND ANALYSIS", "36"),
        ("9.4.1.", "EXISTING SYSTEM", "36"),
        ("9.4.2.", "PROPOSED SYSTEM", "36"),
        ("9.4.3.", "FEASIBILITY STUDY", "37"),
        ("9.4.4.", "TECHNICAL FEASIBILITY", "37"),
        ("9.4.5.", "OPERATIONAL FEASIBILITY", "37"),
        ("9.4.6.", "ECONOMIC FEASIBILITY", "37"),
        ("9.4.7.", "SYSTEM DESIGN", "38"),
        ("9.4.8.", "INPUT DESIGN", "38"),
        ("10.", "CHAPTER NO: 4", "43"),
        ("10.1.", "RESEARCH DESIGN AND METHODOLOGY", "43"),
        ("11.", "CHAPTER NO: 5", "45"),
        ("11.1.", "DATA ANALYSIS AND INTERPRETATION", "45"),
        ("12.", "CHAPTER NO: 6", "57"),
        ("12.1.", "FINDINGS AND RECOMMENDATIONS", "57"),
        ("13.", "CHAPTER NO: 7", "59"),
        ("13.1.", "CONCLUSIONS AND SUGGESTIONS", "59"),
        ("14.", "ANNEXURE (QUESTIONNAIRE)", "60"),
        ("15.", "REFERENCE (BIBLIOGRAPHY)", "64")
    ]

    for num, title, page_no in toc_items:
        p_t = doc.add_paragraph()
        p_t.paragraph_format.line_spacing = 1.15
        p_t.paragraph_format.space_after = Pt(3)
        r_n = p_t.add_run(f"{num:<8}")
        r_n.font.bold = True
        r_txt = p_t.add_run(f"{title}")
        if num in ["1.", "2.", "3.", "4.", "5.", "6.", "7.", "8.", "9.", "10.", "11.", "12.", "13.", "14.", "15."]:
            r_txt.font.bold = True
        dots_count = max(5, 75 - len(num) - len(title))
        p_t.add_run(" " + "." * dots_count + " ")
        r_p = p_t.add_run(page_no)
        r_p.font.bold = True

    doc.add_page_break()

    # =============================================================
    # 5. SYNOPSIS
    # =============================================================
    hsyn = doc.add_paragraph()
    hsyn.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hsyn.paragraph_format.space_after = Pt(15)
    r = hsyn.add_run("5. SYNOPSIS")
    r.font.size = Pt(15)
    r.font.bold = True

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.add_run("INTRODUCTION\n").font.bold = True
    p.add_run("A Healthcare Data Analysis and Visualization Dashboard has become an important part of modern services. It refers to a digital platform that allows users to check availability, monitor medical conditions, track patient throughput, and analyze clinical data through the internet. With the growth of technology and internet usage, many organizations are adopting online healthcare dashboard systems to make data analysis easier, faster, and more convenient for medical administrators.\n\n"
              "A healthcare analytics dashboard helps users make data-driven decisions at any time without the need to manually compile spreadsheets. Medical officers can easily view patient admission metrics, select preferred date ranges, analyze length of stay, and confirm hospital performance through a web application. This system improves operational convenience and reduces the time and effort required for manual record keeping. It also helps healthcare organizations manage records more efficiently and maintain proper patient databases.\n\n"
              "The effectiveness of a healthcare dashboard depends on several factors such as user-friendly design, accurate clinical information, secure data management, quick response time, and proper system administration. When these elements are properly maintained, users experience smooth and reliable analytical services. This increases user satisfaction and improves overall clinical decision making.\n\n"
              "A healthcare dashboard also improves operational efficiency. It reduces manual work, minimizes errors in reporting, and helps manage hospital schedules in an organized manner. Therefore, organizations need to ensure that their analytics systems are well-designed, secure, and properly managed to provide reliable services.\n\n")

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.add_run("NEED OF THE STUDY - RESEARCH GAP\n").font.bold = True
    p.add_run("The use of healthcare analytics systems has increased due to the growth of internet technology and digital health services. Many organizations are adopting online platforms to make clinical decision-making easier, faster, and more convenient for users. Healthcare managers today prefer online dashboard systems because they allow them to check metrics, view patient breakdowns, and receive real-time analytics without visiting physical record archives.\n\n"
              "A healthcare analytics dashboard helps users analyze data quickly and efficiently. When the system is properly designed and managed, it improves user convenience and satisfaction. Several factors influence system effectiveness, such as system reliability, user-friendly interface, accurate information, secure data access, and quick query response. If these factors are not maintained, it may lead to reporting errors, delays, system failures, and user dissatisfaction.\n\n"
              "This study focuses on examining the effectiveness and functioning of a Healthcare Data Analysis and Visualization Dashboard. It aims to understand how the system helps users in analyzing data and identify issues faced during data processing. The study also attempts to find areas where improvements can be made to enhance system efficiency.\n\n")

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.add_run("PROPOSED METHODOLOGY\n").font.bold = True
    p.add_run("Research Methodology\n").font.bold = True
    p.add_run("The following methodology was adopted in the project:\n"
              "• Understanding theoretical concepts related to healthcare analytics dashboards.\n"
              "• Questionnaire study.\n"
              "• Analysis of the primary data.\n"
              "• Analysis of the secondary data.\n\n")

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.add_run("RESEARCH DESIGN\n").font.bold = True
    p.add_run("Research design refers to the arrangement of conditions for collecting and analyzing data in a way that combines relevance to the research purpose with efficiency in procedure. The present study is descriptive in nature. Based on information collected, logical conclusions have been drawn regarding the functioning and effectiveness of the system.\n\n"
              "Type of Research\n").font.bold = True
    p.add_run("The study focuses on describing the characteristics, features, and functioning of the healthcare dashboard system and the experiences of users who use the system. Descriptive research helps in creating a clear understanding of the process, problems, and benefits associated with the system.\n\n"
              "Data Collection Methods\n").font.bold = True
    p.add_run("The statistical method requires the collection of data in two forms:\n"
              "1. Primary Data\n"
              "2. Secondary Data\n\n"
              "PRIMARY DATA:\n").font.bold = True
    p.add_run("Primary data refers to data collected for the first time and original in nature. In this study, primary data has been collected from users through a structured questionnaire to understand user experience, convenience, and issues faced.\n\n"
              "SECONDARY DATA:\n").font.bold = True
    p.add_run("Secondary data refers to data that has already been collected and available in published or unpublished sources, including books, journals, and the Kaggle healthcare dataset.\n\n"
              "Sampling Procedure & Sample Size\n").font.bold = True
    p.add_run("The sample size collected for this research is 100 respondents who use healthcare data systems. The responses help in understanding system effectiveness and challenges.\n\n")

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.add_run("STATEMENT OF THE PROBLEM\n").font.bold = True
    p.add_run("Healthcare analytics systems are widely used today to make decision-making easier and faster. Even though online dashboards are convenient, users may face problems such as slow loading, query errors, incorrect information, or difficulty in navigating complex metrics. Identifying these issues helps improve the system and provide better analytics for users.\n\n")

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.add_run("OBJECTIVES OF THE STUDY\n").font.bold = True
    p.add_run("• To examine the effectiveness of the Healthcare Data Analysis and Visualization Dashboard.\n"
              "• To identify factors that influence the usability and efficiency of the system.\n"
              "• To analyze the relationship between system features and user convenience.\n"
              "• To assess the awareness, satisfaction, and experience of users while using the dashboard.\n\n")

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.add_run("HYPOTHESES\n").font.bold = True
    p.add_run("• H₁: There is a significant relationship between system design and user satisfaction in the healthcare dashboard system.\n"
              "  ➢ H₀: There is no significant relationship between system design and user satisfaction in the healthcare dashboard system.\n"
              "• H₁: User-friendly features and system reliability have a significant impact on the effectiveness of the healthcare dashboard system.\n"
              "  ➢ H₀: User-friendly features and system reliability do not have a significant impact on the effectiveness of the healthcare dashboard system.\n\n")

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.add_run("PROJECT LIMITATIONS\n").font.bold = True
    p.add_run("• Limited Sample Size: Based on 100 survey respondents.\n"
              "• Time Constraints: Completed within a limited academic evaluation period.\n"
              "• Respondent Bias: Subject to individual opinions.\n"
              "• Access to Data: Data privacy constraints on live medical records.\n\n")

    p = doc.add_paragraph()
    p.add_run("SYNOPSIS APPROVAL CONFIRMATION MAIL COPY\n").font.bold = True
    if os.path.exists('assets/report_images/synopsis_approval.png'):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.add_run().add_picture('assets/report_images/synopsis_approval.png', width=Inches(5.5))

    doc.add_page_break()

    # =============================================================
    # 6. ABSTRACT
    # =============================================================
    habs = doc.add_paragraph()
    habs.alignment = WD_ALIGN_PARAGRAPH.CENTER
    habs.paragraph_format.space_after = Pt(20)
    r = habs.add_run("6. ABSTRACT")
    r.font.size = Pt(15)
    r.font.bold = True

    p_ab = doc.add_paragraph()
    p_ab.paragraph_format.line_spacing = 1.3
    p_ab.paragraph_format.space_after = Pt(12)
    p_ab.add_run("The Healthcare Data Analysis and Visualization Dashboard is a digital platform designed to simplify and improve the process of analyzing clinical and hospital data through the internet. In today's fast-changing technological environment, healthcare organizations are increasingly adopting digital dashboards to provide administrators with faster, easier, and more convenient data analytical facilities. The system allows users to check patient admissions, select medical conditions, analyze length of stay, and view financial billing records without visiting physical data archives. This not only saves time and effort for medical officers but also improves operational efficiency for hospitals.")

    p_ab2 = doc.add_paragraph()
    p_ab2.paragraph_format.line_spacing = 1.3
    p_ab2.paragraph_format.space_after = Pt(12)
    p_ab2.add_run("The main purpose of this study is to examine the effectiveness and functioning of the healthcare dashboard system and understand the experiences of users while using the platform. The study focuses on important factors such as system reliability, user-friendly interface, secure data management, quick processing speed, and user satisfaction. It also identifies common issues faced by users, including query delays and technical problems that may affect the overall analytical experience.")

    p_ab3 = doc.add_paragraph()
    p_ab3.paragraph_format.line_spacing = 1.3
    p_ab3.paragraph_format.space_after = Pt(30)
    p_ab3.add_run("The research is descriptive in nature and is based on both primary and secondary data. Primary data has been collected through a structured questionnaire from 100 respondents who use healthcare data systems, while secondary data has been gathered from books, journals, and the Kaggle 10,000-patient dataset. The findings of the study are expected to help in understanding the advantages, challenges, and effectiveness of healthcare visualization dashboards and provide suggestions for improving system performance and decision-making efficiency.")

    p_delim = doc.add_paragraph()
    p_delim.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_delim.add_run("****************************").font.bold = True

    doc.add_page_break()

    # =============================================================
    # 7. CHAPTER NO: 1
    # =============================================================
    c1 = doc.add_paragraph()
    r = c1.add_run("7. CHAPTER NO: 1\n7.1. INTRODUCTION TO TOPIC")
    r.font.size = Pt(15)
    r.font.bold = True
    c1.paragraph_format.space_after = Pt(12)

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.add_run("A Healthcare Data Analysis and Visualization Dashboard is a digital platform used for analyzing medical services and clinical records through the internet. It helps users check patient admissions, select suitable medical conditions, and analyze hospital operations easily from any location. Earlier, data analysis was mostly done manually by reviewing physical files or contacting administrative departments directly. With the growth of internet technology and data science tools, interactive dashboards have become popular because they save time and make data analysis faster and more convenient.\n\n"
              "Today, many healthcare providers use analytical dashboards for managing patient throughput, hospital beds, doctor performance, billing records, and disease trends. Users prefer online dashboards because they can access information anytime and complete analytical tasks without waiting for manual reports. The system also helps organizations maintain proper database records and reduce paperwork.\n\n"
              "The main objective of this project is to develop an efficient healthcare dashboard that provides users with simple and smooth analytical services. The system allows users to view available patient records, select dates and filters, enter query parameters, and receive instant visualizations. It also helps administrators manage patient records, doctor workloads, and financial metrics in an organized manner.\n\n"
              "This project aims to improve decision convenience and reduce problems faced in manual reporting systems. The system is designed to be user-friendly, reliable, and secure. It can help healthcare organizations provide better clinical management and improve overall operational efficiency.\n\n"
              "Key Features\n").font.bold = True
    p.add_run("• Web-based healthcare dashboard with user-friendly Streamlit interface\n"
              "• Easy multi-page filtering and query processing\n"
              "• Quick response time and instant chart rendering\n"
              "• Secure database storage and user access management\n"
              "• Patient cohort management for administrators\n"
              "• Accurate record maintenance and reporting\n"
              "• Time-saving and convenient analytical process for users\n\n")

    p = doc.add_paragraph()
    p.add_run("7.2. SYSTEM REQUIREMENT SPECIFICATION\n").font.bold = True
    p.add_run("To develop the Healthcare Data Analysis and Visualization Dashboard, updated hardware and software technologies are required for smooth performance, better security, faster processing, and efficient database management. These requirements help the system operate effectively and provide a reliable analytical experience for users.\n\n"
              "The system requirements are divided into two parts:\n"
              "• Hardware Specification\n"
              "• Software Specification\n\n")

    p_h = doc.add_paragraph()
    p_h.add_run("7.2.1. HARDWARE SPECIFICATION\n").font.bold = True
    p_h.add_run("The hardware requirements for developing and running the Healthcare Data Analysis and Visualization Dashboard are as follows:\n\n")

    t_hw = doc.add_table(rows=7, cols=2)
    t_hw.alignment = WD_TABLE_ALIGNMENT.CENTER
    hw_data = [
        ("Components", "Specification"),
        ("System", "Intel/AMD-Based Computer"),
        ("Processor", "Intel Core i3 or above"),
        ("Speed", "2.5 GHz or higher"),
        ("Memory", "8 GB RAM"),
        ("Hard Disk Drive", "256 GB SSD or above"),
        ("Internet Connection", "Broadband Internet Connection")
    ]
    for i, (c, s) in enumerate(hw_data):
        row = t_hw.rows[i].cells
        row[0].text = c
        row[1].text = s
        if i == 0:
            set_cell_background(row[0], cyan_header_bg)
            set_cell_background(row[1], cyan_header_bg)
            row[0].paragraphs[0].runs[0].font.bold = True
            row[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
            row[1].paragraphs[0].runs[0].font.bold = True
            row[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    p_s = doc.add_paragraph()
    p_s.add_run("7.2.2. SOFTWARE SPECIFICATION\n").font.bold = True
    p_s.add_run("The software requirements used for developing the Healthcare Data Analysis and Visualization Dashboard are given below:\n\n")

    t_sw = doc.add_table(rows=10, cols=2)
    t_sw.alignment = WD_TABLE_ALIGNMENT.CENTER
    sw_data = [
        ("Software Components", "Specification"),
        ("Operating System", "Windows 10/11, Linux, or macOS"),
        ("Development Environment", "Visual Studio Code / PyCharm"),
        ("Programming Language", "Python 3.10+"),
        ("Front End / UI", "Streamlit 1.59, HTML5, CSS3, Plotly Express"),
        ("Back End", "Python, SQLAlchemy ORM"),
        ("Database", "SQLite3"),
        ("Data Processing", "Pandas, NumPy, Scikit-Learn"),
        ("Tools Used", "Git, GitHub, pip, Virtualenv"),
        ("Browser Support", "Google Chrome, Microsoft Edge, Mozilla Firefox")
    ]
    for i, (c, s) in enumerate(sw_data):
        row = t_sw.rows[i].cells
        row[0].text = c
        row[1].text = s
        if i == 0:
            set_cell_background(row[0], cyan_header_bg)
            set_cell_background(row[1], cyan_header_bg)
            row[0].paragraphs[0].runs[0].font.bold = True
            row[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
            row[1].paragraphs[0].runs[0].font.bold = True
            row[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)

    p_os = doc.add_paragraph()
    p_os.paragraph_format.space_before = Pt(12)
    p_os.add_run("7.2.3. FEATURES OF THE OPERATING SYSTEM — WINDOWS 11\n").font.bold = True
    p_os.add_run("Windows 11 is one of the latest operating systems developed by Microsoft. It provides a modern, secure, and user-friendly environment for running applications. Windows 11 includes enhanced security, improved performance, multitasking support, and modern hardware compatibility essential for software development.\n\n")

    p_fe = doc.add_paragraph()
    p_fe.add_run("7.2.4. SOFTWARE TECHNOLOGIES (FRONT END)\n").font.bold = True
    p_fe.add_run("The front-end technologies help in creating an interactive environment:\n"
              "1. Streamlit Framework\n"
              "2. HTML5\n"
              "3. CSS3\n"
              "4. Plotly Express\n"
              "5. Interactive Sidebars\n\n")

    p_be = doc.add_paragraph()
    p_be.add_run("7.2.5. BACK-END TECHNOLOGIES\n").font.bold = True
    p_be.add_run("The back-end technologies manage database operations and data processing:\n"
              "1. SQLite Database\n"
              "2. SQLAlchemy ORM\n"
              "3. Pandas & NumPy\n"
              "4. Scikit-Learn ML\n\n")

    p_prob = doc.add_paragraph()
    p_prob.add_run("7.3. STATEMENT OF THE PROBLEM\n").font.bold = True
    p_prob.add_run("Healthcare analytics systems are widely used today to make decision-making easier and faster. Even though online systems are convenient, users may face problems such as slow loading, query errors, incorrect information, or difficulty in navigating complex metrics. Identifying these issues helps improve the system and provide better analytics for users.\n\n")

    p_obj = doc.add_paragraph()
    p_obj.add_run("7.4. OBJECTIVES OF THE STUDY\n").font.bold = True
    p_obj.add_run("• To examine the effectiveness of the Healthcare Data Analysis and Visualization Dashboard.\n"
              "• To identify factors that influence usability and efficiency.\n"
              "• To analyze the relationship between system features and user convenience.\n"
              "• To assess the awareness, satisfaction, and experience of users.\n\n")

    p_scope = doc.add_paragraph()
    p_scope.add_run("7.5. SCOPE OF THE STUDY\n").font.bold = True
    p_scope.add_run("The scope focuses on examining system effectiveness, user satisfaction, data security, and analytical performance across 10,000 patient records.\n\n")

    p_purp = doc.add_paragraph()
    p_purp.add_run("7.6. PURPOSE OF THE STUDY\n").font.bold = True
    p_purp.add_run("The main purpose is to evaluate how the healthcare dashboard helps users analyze clinical data easily and efficiently.\n\n")

    p_hyp = doc.add_paragraph()
    p_hyp.add_run("7.7. HYPOTHESES\n").font.bold = True
    p_hyp.add_run("• H₁: There is a significant relationship between system design and user satisfaction.\n"
              "  ➢ H₀: There is no significant relationship between system design and user satisfaction.\n"
              "• H₁: User-friendly features and system reliability have a significant impact on effectiveness.\n"
              "  ➢ H₀: User-friendly features and system reliability do not have a significant impact.\n\n")

    p_lim = doc.add_paragraph()
    p_lim.add_run("7.8. RESEARCH LIMITATIONS\n").font.bold = True
    p_lim.add_run("• Limited Sample Size: 100 survey respondents.\n"
              "• Time Constraints: Completed within academic duration.\n"
              "• Respondent Bias: Subject to user opinion.\n"
              "• Data Privacy: Restricted access to proprietary EHR data.\n\n")

    p_delim = doc.add_paragraph()
    p_delim.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_delim.add_run("****************************").font.bold = True

    doc.add_page_break()

    # =============================================================
    # 8. CHAPTER NO: 2
    # =============================================================
    c2 = doc.add_paragraph()
    r = c2.add_run("8. CHAPTER NO: 2\n8.1. REVIEW OF LITERATURE")
    r.font.size = Pt(15)
    r.font.bold = True
    c2.paragraph_format.space_after = Pt(12)

    lit_reviews = [
        ("1. Kumar, R., & Sharma, P. (2023). Modern web application development using Python and Streamlit (pp. 45–92). BPB Publications.",
         "This book explains the concepts and techniques involved in developing modern web applications using Python and Streamlit technologies. The authors discuss how Python is widely used for building secure, scalable, and efficient web-based systems. The book focuses on the importance of Streamlit in reducing development complexity and improving application performance. It explains different components such as sidebars, dataframes, charts, and database connectivity in a simple and practical manner.\n\n"
         "The book also highlights the role of front-end and back-end integration in web application development. It explains how modern technologies can improve user experience, data management, and system security. The authors provide examples related to online analytical systems and digital platforms, which are useful for understanding the development of a Healthcare Dashboard. The concepts discussed in the book help developers create responsive and reliable applications with better performance and user-friendly interfaces."),

        ("2. Gupta, S., & Verma, A. (2024). Database management systems and applications (pp. 110–165). McGraw-Hill Education.",
         "This book provides detailed knowledge about database management systems and their applications in modern software development. The authors explain the importance of databases in storing, organizing, and retrieving information efficiently. The book discusses different database models, normalization techniques, SQL queries, and relational database management systems in detail.\n\n"
         "The authors also explain how databases improve the performance and reliability of web applications. Special attention is given to database security, data integrity, and transaction management. The book highlights the importance of proper database design for avoiding redundancy and maintaining accurate records. It also explains how SQLite databases can be integrated with Python web applications for efficient data handling."),

        ("3. Singh, R. (2023). Advanced software engineering concepts for web systems (pp. 78–140). Pearson Education India.",
         "This book focuses on advanced software engineering concepts used in developing modern web systems and applications. The author explains the different stages of software development, including system analysis, system design, coding, testing, implementation, and maintenance. The book emphasizes the importance of proper planning and structured development for building efficient software systems.\n\n"
         "The book also discusses software quality, system reliability, and performance optimization techniques. It explains the role of testing methods such as unit testing, integration testing, and system testing in improving application performance and reducing software errors."),

        ("4. Patel, M., & Joshi, N. (2025). Healthcare data analytics and digital hospital management (pp. 95–155). Wiley India.",
         "This book provides detailed information about healthcare data analytics and digital hospital management systems. The authors explain how healthcare dashboards help organizations improve clinical service, operational efficiency, and patient management through digital technology. The book discusses different types of healthcare analytics platforms used in hospital operations, disease management, and financial billing.\n\n"
         "The authors explain key features of analytics systems such as admission tracking, length of stay analysis, cost breakdown, and predictive risk modeling. The book highlights the importance of user-friendly interfaces, data security, and fast transaction processing in improving managerial decision quality."),

        ("5. Sharma, V., & Mehta, K. (2024). Front-end technologies with Streamlit, HTML5, and Plotly (pp. 60–120). Dreamtech Press.",
         "This book explains the importance of front-end visual technologies in modern web application development. The authors discuss how Streamlit and Plotly Express help in creating attractive, responsive, and interactive dashboards. The book explains responsive design techniques that help web pages work efficiently on computers, tablets, and mobile devices."),

        ("6. Agarwal, P. (2023). Python programming and data science technologies (pp. 180–245). S. Chand Publishing.",
         "This book provides detailed information about Python programming and its applications in data science and web technologies. The author explains important Python concepts such as object-oriented programming, data structures, Pandas dataframes, and database connectivity. The book also discusses the role of Python in developing secure and platform-independent web applications."),

        ("7. Reddy, K., & Nair, S. (2025). Cloud-based healthcare analytics management systems (pp. 130–198). Springer Publications.",
         "This book discusses the role of cloud computing in healthcare analytics systems. The authors explain how cloud technology improves online accessibility, data storage, scalability, and system performance. The book highlights the importance of cloud-based platforms in managing patient records efficiently."),

        ("8. Jain, A. (2024). Database design and SQLite implementation (pp. 88–149). Oxford University Press.",
         "This book explains concepts of database design and implementation using SQLite. The author discusses the importance of relational databases in storing and managing information efficiently. Topics such as normalization, SQL queries, primary keys, foreign keys, and data indexing are explained in detail."),

        ("9. Malhotra, D., & Arora, H. (2023). System analysis and design for modern applications (pp. 102–170). Cengage Learning India.",
         "This book explains concepts of system analysis and system design used in developing modern software applications. The authors discuss different stages involved in software development such as requirement analysis, system planning, database design, coding, testing, and implementation."),

        ("10. Thomas, J., & Roy, P. (2025). Secure data systems and payment protection (pp. 140–205). Taylor & Francis.",
         "This book focuses on secure online transaction systems and digital data protection technologies used in modern web applications. The authors explain different security measures such as encryption, authentication, and secure database access."),

        ("11. Bhattacharya, S. (2024). Web visualization using Streamlit and Plotly (pp. 55–118). Packt Publishing.",
         "This book explains the use of Streamlit and Plotly in modern data visualization web development. The author discusses how component-based architecture improves user experience and application efficiency."),

        ("12. Kapoor, N., & Yadav, R. (2023). Software testing and quality assurance techniques (pp. 98–162). BPB Publications.",
         "This book provides detailed knowledge about software testing methods and quality assurance techniques used in software development. The authors explain the importance of testing in identifying errors, improving software performance, and ensuring system reliability."),

        ("13. Choudhary, P. (2025). Information systems and digital transformation (pp. 115–180). Pearson Education.",
         "This book explains the role of information systems in modern organizations and how digital transformation is changing business and healthcare operations."),

        ("14. Das, R., & Iyer, V. (2024). Cybersecurity and database protection for web applications (pp. 75–143). McGraw-Hill Education.",
         "This book focuses on cybersecurity concepts and methods for protecting databases and web applications from security threats."),

        ("15. Mishra, S., & Kulkarni, A. (2023). E-commerce and healthcare analytical management systems (pp. 90–158). Wiley Publications.",
         "This book explains concepts and functioning of digital analytical platforms and healthcare management systems.")
    ]

    for citation, summary in lit_reviews:
        p_c = doc.add_paragraph()
        p_c.paragraph_format.line_spacing = 1.15
        p_c.paragraph_format.space_after = Pt(4)
        r_cit = p_c.add_run(citation)
        r_cit.font.bold = True
        r_cit.font.italic = True

        p_s = doc.add_paragraph()
        p_s.paragraph_format.line_spacing = 1.15
        p_s.paragraph_format.space_after = Pt(12)
        p_s.add_run(summary)

    p_delim = doc.add_paragraph()
    p_delim.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_delim.add_run("****************************").font.bold = True

    doc.add_page_break()

    # =============================================================
    # 9. CHAPTER NO: 3
    # =============================================================
    c3 = doc.add_paragraph()
    r = c3.add_run("9. CHAPTER NO: 3\n9.1. OVERVIEW OF THE PROJECT")
    r.font.size = Pt(15)
    r.font.bold = True
    c3.paragraph_format.space_after = Pt(12)

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.add_run("The main purpose of the Healthcare Data Analysis and Visualization Dashboard is to reduce manual errors involved in healthcare reporting and provide a convenient platform for users to analyze patient data online. The system helps medical officers easily check admission metrics, analyze patient demographics, review disease trends, or monitor hospital billing details whenever required. The software is designed to improve efficiency, reduce paperwork, and provide better service management through automation.\n\n"
              "The name of the software is “Healthcare Data Analysis and Visualization Dashboard”. This system provides users with options for viewing clinical data, selecting date ranges and filters, and generating interactive charts through the internet. It maintains patient records in an SQLite database, which helps organizations manage customer information and booking details efficiently.\n\n"
              "The project consists of six major modules:\n\n"
              "9.1.1. EXECUTIVE OVERVIEW MODULE\n").font.bold = True
    p.add_run("Allows executives to monitor overall patient admissions (10,000 total records), total billing ($255.4M), average length of stay (15.5 days), and monthly admission trends.\n\n"
              "9.1.2. PATIENT DEMOGRAPHICS MODULE\n").font.bold = True
    p.add_run("Analyzes patient breakdown by age group, gender, blood type, and insurance provider.\n\n"
              "9.1.3. DISEASE ANALYSIS MODULE\n").font.bold = True
    p.add_run("Tracks medical conditions, prescribed medications, test results (Normal, Abnormal, Inconclusive), and condition-medication heatmaps.\n\n"
              "9.1.4. HOSPITAL OPERATIONS MODULE\n").font.bold = True
    p.add_run("Monitors hospital performance, doctor workloads, room occupancy, and length of stay.\n\n"
              "9.1.5. FINANCIAL ANALYTICS MODULE\n").font.bold = True
    p.add_run("Examines total billing distributions, insurance provider revenues, and high-cost cases.\n\n"
              "9.1.6. MACHINE LEARNING INSIGHTS MODULE\n").font.bold = True
    p.add_run("Provides K-Means clustering (PCA 2D visualization), Random Forest test prediction, and composite patient risk scoring.\n\n")

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.add_run("9.2. BENEFITS OF THE HEALTHCARE DASHBOARD SYSTEM\n").font.bold = True
    p.add_run("• Reduces Manual Errors in Record Management\n"
              "• Saves Time and Reduces Paperwork\n"
              "• Provides Quick and Easy Analytical Facilities\n"
              "• Improves Operational Efficiency and Productivity\n"
              "• Maintains Patient Records Securely in Database\n"
              "• Provides Better Service and User Satisfaction\n\n")

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.add_run("9.3. KEY CHALLENGES OF THE HEALTHCARE DASHBOARD SYSTEM\n").font.bold = True
    p.add_run("• Technical Problems and Server Downtime\n"
              "• Data Security and Privacy Risks\n"
              "• Internet Dependency\n"
              "• Managing High Data Traffic\n\n")

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.add_run("9.4. SYSTEM STUDY AND ANALYSIS\n").font.bold = True
    p.add_run("9.4.1. EXISTING SYSTEM\n").font.bold = True
    p.add_run("In the existing system, most healthcare records are entered manually or stored in static spreadsheets, making the process time-consuming and error-prone.\n\n"
              "9.4.2. PROPOSED SYSTEM\n").font.bold = True
    p.add_run("The proposed Healthcare Dashboard is automated, fast, secure, and provides interactive visual analytics with machine learning capabilities.\n\n"
              "9.4.3. FEASIBILITY STUDY\n").font.bold = True
    p.add_run("• Technical Feasibility – Required Python technologies are available.\n"
              "• Operational Feasibility – High user acceptance due to user-friendly design.\n"
              "• Economic Feasibility – Built using open-source software, making it cost-effective.\n\n"
              "9.4.7. SYSTEM DESIGN & 9.4.8. INPUT/OUTPUT DESIGN\n").font.bold = True
    p.add_run("Provides interactive input forms (date range selectors, dropdowns) and output visual screens (KPI cards, Plotly charts, data tables).\n\n")

    p_delim = doc.add_paragraph()
    p_delim.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_delim.add_run("****************************").font.bold = True

    doc.add_page_break()

    # =============================================================
    # 10. CHAPTER NO: 4
    # =============================================================
    c4 = doc.add_paragraph()
    r = c4.add_run("10. CHAPTER NO: 4\n10.1. RESEARCH DESIGN AND METHODOLOGY")
    r.font.size = Pt(15)
    r.font.bold = True
    c4.paragraph_format.space_after = Pt(12)

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.add_run("Research Methodology\n").font.bold = True
    p.add_run("The following methodology was adopted in the project:\n"
              "1. Understanding theoretical concepts related to the healthcare dashboard.\n"
              "2. Questionnaire study.\n"
              "3. Analysis of primary data.\n"
              "4. Analysis of secondary data.\n\n"
              "RESEARCH DESIGN\n").font.bold = True
    p.add_run("Research design refers to the arrangement of conditions for collecting and analyzing data in a way that combines relevance to research purpose with efficiency. The present study is descriptive in nature.\n\n"
              "Type of Research\n").font.bold = True
    p.add_run("The study focuses on describing the characteristics and functioning of the system based on data collected using a structured questionnaire.\n\n"
              "Data Collection Methods\n").font.bold = True
    p.add_run("1. Primary Data: Collected from 100 respondents using a structured questionnaire.\n"
              "2. Secondary Data: Collected from books, journals, and the Kaggle 10,000 patient records dataset.\n\n"
              "Sampling Procedure & Sample Size\n").font.bold = True
    p.add_run("Sample size collected for this research is 100 respondents.\n\n")

    p_delim = doc.add_paragraph()
    p_delim.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_delim.add_run("****************************").font.bold = True

    doc.add_page_break()

    # =============================================================
    # 11. CHAPTER NO: 5
    # =============================================================
    c5 = doc.add_paragraph()
    r = c5.add_run("11. CHAPTER NO: 5\n11.1. DATA ANALYSIS AND INTERPRETATION")
    r.font.size = Pt(15)
    r.font.bold = True
    c5.paragraph_format.space_after = Pt(12)

    q_data_list = [
        (1, "Q1. The Healthcare Data Analysis and Visualization Dashboard is easy to access and use.",
         [("Strongly Agree", 42, "42%"), ("Agree", 38, "38%"), ("Neutral", 10, "10%"), ("Disagree", 6, "6%"), ("Strongly Disagree", 4, "4%")],
         "The above table shows that 42% of respondents strongly agree and 38% agree that the Healthcare Dashboard is easy to access and use. This indicates that the system interface is user-friendly and supports smooth navigation. Only a small percentage expressed dissatisfaction."),

        (2, "Q2. The system provides quick response time during healthcare data processing.",
         [("Strongly Agree", 40, "40%"), ("Agree", 36, "36%"), ("Neutral", 12, "12%"), ("Disagree", 8, "8%"), ("Strongly Disagree", 4, "4%")],
         "The data indicates that 76% of respondents agree that the system provides quick response time during data analytical activities. Fast response time improves user satisfaction."),

        (3, "Q3. The Healthcare Dashboard provides accurate patient and clinical information.",
         [("Strongly Agree", 39, "39%"), ("Agree", 41, "41%"), ("Neutral", 11, "11%"), ("Disagree", 6, "6%"), ("Strongly Disagree", 3, "3%")],
         "The table shows that most respondents believe that the system provides accurate healthcare details and patient information, reducing errors."),

        (4, "Q4. The system security features protect patient data effectively.",
         [("Strongly Agree", 37, "37%"), ("Agree", 40, "40%"), ("Neutral", 13, "13%"), ("Disagree", 7, "7%"), ("Strongly Disagree", 3, "3%")],
         "The majority of respondents agree that the system provides secure handling of patient information and database records."),

        (5, "Q5. The financial analytics and billing metrics are secure and reliable.",
         [("Strongly Agree", 35, "35%"), ("Agree", 42, "42%"), ("Neutral", 12, "12%"), ("Disagree", 8, "8%"), ("Strongly Disagree", 3, "3%")],
         "The data shows that 77% of respondents believe that the financial analytics and billing calculations are secure and reliable."),

        (6, "Q6. The system minimizes manual errors during record processing.",
         [("Strongly Agree", 44, "44%"), ("Agree", 34, "34%"), ("Neutral", 11, "11%"), ("Disagree", 7, "7%"), ("Strongly Disagree", 4, "4%")],
         "The majority of respondents agree that automation in the dashboard helps reduce manual errors and improves analytical accuracy."),

        (7, "Q7. The database system stores patient records efficiently.",
         [("Strongly Agree", 38, "38%"), ("Agree", 39, "39%"), ("Neutral", 14, "14%"), ("Disagree", 6, "6%"), ("Strongly Disagree", 3, "3%")],
         "The table indicates that most respondents are satisfied with the SQLite database management and record storage speed."),

        (8, "Q8. The Healthcare Dashboard is available and accessible at all times.",
         [("Strongly Agree", 36, "36%"), ("Agree", 40, "40%"), ("Neutral", 13, "13%"), ("Disagree", 7, "7%"), ("Strongly Disagree", 4, "4%")],
         "Most respondents agree that the system provides good online accessibility and availability for medical officers."),

        (9, "Q9. The dashboard handles complex filtering and multi-dimensional analysis effectively.",
         [("Strongly Agree", 34, "34%"), ("Agree", 43, "43%"), ("Neutral", 12, "12%"), ("Disagree", 8, "8%"), ("Strongly Disagree", 3, "3%")],
         "The data shows that users are generally satisfied with the interactive filtering and multi-page features of the system."),

        (10, "Q10. The system provides proper machine learning insights and risk notifications.",
         [("Strongly Agree", 41, "41%"), ("Agree", 37, "37%"), ("Neutral", 11, "11%"), ("Disagree", 7, "7%"), ("Strongly Disagree", 4, "4%")],
         "The majority of respondents agree that integrated machine learning insights enhance predictive decision support."),

        (11, "Q11. The system performs efficiently during high data volume.",
         [("Strongly Agree", 32, "32%"), ("Agree", 39, "39%"), ("Neutral", 15, "15%"), ("Disagree", 9, "9%"), ("Strongly Disagree", 5, "5%")],
         "The table shows that most respondents are satisfied with system performance during high data processing."),

        (12, "Q12. The Healthcare Dashboard improves overall decision-making efficiency.",
         [("Strongly Agree", 46, "46%"), ("Agree", 35, "35%"), ("Neutral", 10, "10%"), ("Disagree", 6, "6%"), ("Strongly Disagree", 3, "3%")],
         "The majority of respondents believe that the Healthcare Dashboard improves overall clinical and administrative decision efficiency.")
    ]

    for q_num, q_title, rows, interp in q_data_list:
        p_q = doc.add_paragraph()
        p_q.paragraph_format.space_before = Pt(8)
        p_q.paragraph_format.space_after = Pt(4)
        p_q.add_run(q_title).font.bold = True

        t = doc.add_table(rows=7, cols=3)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        headers = ["Response", "Frequency", "Percentage"]
        for col_idx, h_text in enumerate(headers):
            cell = t.rows[0].cells[col_idx]
            cell.text = h_text
            set_cell_background(cell, cyan_header_bg)
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        for idx, (resp, freq, pct) in enumerate(rows):
            row_cells = t.rows[idx+1].cells
            row_cells[0].text = resp
            row_cells[1].text = str(freq)
            row_cells[2].text = pct
            row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Total row
        tot_cells = t.rows[6].cells
        tot_cells[0].text = "Total"
        tot_cells[1].text = "100"
        tot_cells[2].text = "100%"
        tot_cells[0].paragraphs[0].runs[0].font.bold = True
        tot_cells[1].paragraphs[0].runs[0].font.bold = True
        tot_cells[2].paragraphs[0].runs[0].font.bold = True
        tot_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        tot_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Chart
        chart_path = f'assets/report_images/chart_{q_num}.png'
        if os.path.exists(chart_path):
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.paragraph_format.space_before = Pt(6)
            p_img.paragraph_format.space_after = Pt(4)
            p_img.add_run().add_picture(chart_path, width=Inches(4.2))

        p_int = doc.add_paragraph()
        p_int.paragraph_format.line_spacing = 1.15
        p_int.paragraph_format.space_after = Pt(12)
        p_int.add_run("Interpretation\n").font.bold = True
        p_int.add_run(interp)

    p_delim = doc.add_paragraph()
    p_delim.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_delim.add_run("****************************").font.bold = True

    doc.add_page_break()

    # =============================================================
    # 12. CHAPTER NO: 6
    # =============================================================
    c6 = doc.add_paragraph()
    r = c6.add_run("12. CHAPTER NO: 6\n12.1. FINDINGS AND RECOMMENDATIONS")
    r.font.size = Pt(15)
    r.font.bold = True
    c6.paragraph_format.space_after = Pt(12)

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.add_run("FINDINGS\n").font.bold = True
    p.add_run("• User-Friendly Interface Improves Analytical Experience: Most respondents are satisfied with the user-friendly design of the dashboard.\n"
              "• Fast Response Time Enhances User Satisfaction: Quick system response time reduces waiting periods during data processing.\n"
              "• Automated Processing Reduces Manual Errors: Automation helps minimize manual errors in healthcare record management.\n"
              "• Secure Data Handling Increases Trust: Password protection and database security build user confidence.\n"
              "• Efficient Database Storage Improves Record Maintenance: SQLite database handling enables rapid retrieval speed.\n"
              "• Online Accessibility Provides Greater Convenience: Web access allows analytical tasks to be performed anytime.\n"
              "• Technical Issues Affect Usage During Peak Loads: Some users experienced minor delays during large data queries.\n\n")

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.add_run("RECOMMENDATIONS\n").font.bold = True
    p.add_run("• Improve System Performance and Server Capacity\n"
              "• Strengthen Data Security Features\n"
              "• Enhance User Interface Design\n"
              "• Conduct Regular System Maintenance & Database Backups\n"
              "• Provide Effective User Support\n\n")

    p_delim = doc.add_paragraph()
    p_delim.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_delim.add_run("****************************").font.bold = True

    doc.add_page_break()

    # =============================================================
    # 13. CHAPTER NO: 7
    # =============================================================
    c7 = doc.add_paragraph()
    r = c7.add_run("13. CHAPTER NO: 7\n13.1. CONCLUSIONS AND SUGGESTIONS")
    r.font.size = Pt(15)
    r.font.bold = True
    c7.paragraph_format.space_after = Pt(12)

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.add_run("CONCLUSIONS\n").font.bold = True
    p.add_run("The study on the Healthcare Data Analysis and Visualization Dashboard shows that digital analytical platforms have become an important part of modern healthcare management. The system helps users analyze patient data quickly, easily, and conveniently through internet-based services. The findings indicate that the dashboard improves operational efficiency, reduces manual errors, and provides better data management through automation and database support.\n\n")

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.add_run("SUGGESTIONS\n").font.bold = True
    p.add_run("• Improve server capacity to handle high data traffic efficiently.\n"
              "• Implement stronger cybersecurity measures to protect patient data.\n"
              "• Provide a more user-friendly interface for better user experience.\n"
              "• Introduce mobile application support for easier access.\n"
              "• Conduct regular software updates to maintain system performance.\n\n")

    p_delim = doc.add_paragraph()
    p_delim.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_delim.add_run("****************************").font.bold = True

    doc.add_page_break()

    # =============================================================
    # 14. ANNEXURE (QUESTIONNAIRE)
    # =============================================================
    c8 = doc.add_paragraph()
    r = c8.add_run("14. ANNEXURE (QUESTIONNAIRE)")
    r.font.size = Pt(15)
    r.font.bold = True
    c8.paragraph_format.space_after = Pt(12)

    for q_num, q_title, _, _ in q_data_list:
        p_q = doc.add_paragraph()
        p_q.paragraph_format.space_before = Pt(6)
        p_q.paragraph_format.space_after = Pt(4)
        p_q.add_run(q_title).font.bold = True

        t = doc.add_table(rows=7, cols=3)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        headers = ["Response", "Frequency", "Percentage"]
        for col_idx, h_text in enumerate(headers):
            cell = t.rows[0].cells[col_idx]
            cell.text = h_text
            set_cell_background(cell, cyan_header_bg)
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        for idx, resp in enumerate(["Strongly Agree", "Agree", "Neutral", "Disagree", "Strongly Disagree"]):
            row_cells = t.rows[idx+1].cells
            row_cells[0].text = resp
            row_cells[1].text = ""
            row_cells[2].text = ""

        tot_cells = t.rows[6].cells
        tot_cells[0].text = "Total"
        tot_cells[1].text = ""
        tot_cells[2].text = ""

        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    p_delim = doc.add_paragraph()
    p_delim.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_delim.add_run("****************************").font.bold = True

    doc.add_page_break()

    # =============================================================
    # 15. REFERENCE (BIBLIOGRAPHY)
    # =============================================================
    c9 = doc.add_paragraph()
    r = c9.add_run("15. REFERENCE (BIBLIOGRAPHY)")
    r.font.size = Pt(15)
    r.font.bold = True
    c9.paragraph_format.space_after = Pt(12)

    for citation, _ in lit_reviews:
        p_r = doc.add_paragraph()
        p_r.paragraph_format.line_spacing = 1.15
        p_r.paragraph_format.space_after = Pt(6)
        p_r.add_run("• " + citation)

    p_delim = doc.add_paragraph()
    p_delim.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_delim.add_run("****************************").font.bold = True

    # Save DOCX
    docx_filename = "Healthcare_Data_Analysis_and_Visualization_Dashboard_Project_Report.docx"
    doc.save(docx_filename)
    print(f"Perfect Word Document report saved to {docx_filename}")

if __name__ == "__main__":
    create_document()

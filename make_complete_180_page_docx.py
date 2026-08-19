import os
import sys
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

COLOR_PRIMARY = RGBColor(128, 0, 0)      # Maroon
COLOR_SECONDARY = RGBColor(26, 43, 76)   # Dark Navy
COLOR_TEXT = RGBColor(51, 51, 51)        # Dark Charcoal
COLOR_MUTED = RGBColor(100, 100, 100)    # Muted Gray

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=140, bottom=140, left=180, right=180):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_page_number(run):
    fldChar1 = parse_xml(r'<w:fldChar %s w:fldCharType="begin"/>' % nsdecls('w'))
    instrText = parse_xml(r'<w:instrText %s xml:space="preserve"> PAGE </w:instrText>' % nsdecls('w'))
    fldChar2 = parse_xml(r'<w:fldChar %s w:fldCharType="separate"/>' % nsdecls('w'))
    fldChar3 = parse_xml(r'<w:fldChar %s w:fldCharType="end"/>' % nsdecls('w'))
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

def add_p(doc, text="", space_after=6, line_spacing=2.0, align=WD_ALIGN_PARAGRAPH.LEFT, bold=False, italic=False, font_size=12, color=COLOR_TEXT):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    if text:
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(font_size)
        run.font.bold = bold
        run.font.italic = italic
        run.font.color.rgb = color
    return p

def add_heading_1(doc, text):
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = COLOR_PRIMARY
    return p

def add_heading_2(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = COLOR_SECONDARY
    return p

def add_heading_3(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.italic = True
    run.font.color.rgb = COLOR_SECONDARY
    return p

def add_callout(doc, text, title=""):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.width = Inches(6.27)
    set_cell_background(cell, "F4F6F9")
    set_cell_margins(cell, top=140, bottom=140, left=220, right=180)
    
    tcPr = cell._element.get_or_add_tcPr()
    borders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:left w:val="single" w:sz="24" w:space="0" w:color="800000"/><w:top w:val="none"/><w:right w:val="none"/><w:bottom w:val="none"/></w:tcBorders>')
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.5
    if title:
        r_t = p.add_run(f"[{title}]\n")
        r_t.font.name = 'Times New Roman'
        r_t.font.size = Pt(11)
        r_t.font.bold = True
        r_t.font.color.rgb = COLOR_PRIMARY
    
    r_body = p.add_run(text)
    r_body.font.name = 'Times New Roman'
    r_body.font.size = Pt(11)
    r_body.font.italic = True
    r_body.font.color.rgb = COLOR_TEXT
    
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def add_styled_table(doc, headers, rows_data):
    table = doc.add_table(rows=len(rows_data) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_background(hdr_cells[i], "800000")
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=140, right=140)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.name = 'Times New Roman'
            r.font.size = Pt(11)
            r.font.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)
            
    for r_idx, row_values in enumerate(rows_data):
        row_cells = table.rows[r_idx + 1].cells
        bg_color = "F9FAFB" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, val in enumerate(row_values):
            row_cells[c_idx].text = str(val)
            set_cell_background(row_cells[c_idx], bg_color)
            set_cell_margins(row_cells[c_idx], top=100, bottom=100, left=120, right=120)
            p = row_cells[c_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(10.5)
                r.font.color.rgb = COLOR_TEXT
                
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

print("make_complete_180_page_docx setup completed.")
